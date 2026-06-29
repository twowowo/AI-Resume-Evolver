"""
AI-Resume-Evolver v3.0 FastAPI 路由网关

端点:
  GET  /health                           — 健康检查
  POST /api/v1/resume/optimize           — 简历优化（SSE 流式事件推送）
  POST /api/v1/resume/chat               — 多轮对话交互编辑（SSE 流式）

一键生成模式 (ONE_CLICK) — SSE 流式四帧推送:
  Frame 1 (radar_init):   PreEvaluator 完成 → 原始简历 6-3-1 雷达指标，前端渲染初始雷达图
  Frame 2 (resume_stream): Editor 完成 → 精修简历全文，前端打字机渐进显示
  Frame 3 (final):        Evaluator+Interviewer 完成 → 终评雷达 + 3 道面试压测题
  Frame 4 (done):         流结束信号

交互模式 (INTERACTIVE) — SSE 流式三帧推送:
  Frame 1 (status):       chat_editor 增量编辑完成 → 当前节点状态 (复用 resume_stream 帧携带 node_status)
  Frame 2 (resume_stream): 更新后的简历全文 + node_status 状态帧
  Frame 3 (final):        Evaluator+Interviewer 完成 → 终评雷达 + 压测题
  Frame N (done):         流结束信号

v3.0 更新:
  - 引入 MemorySaver 断点续传，全图编译挂载 checkpointer
  - 新增 /api/v1/resume/chat 多轮对话端点，支持 thread_id 会话隔离
  - graph 入口 router 自动分流一键模式 / 交互模式
  - ONE_CLICK 端点对接共享图实例，自动生成 session_id 供后续 chat 续接
"""
#uvicorn main:app --host 127.0.0.1 --port 8001 --reload
#npm run dev
#taskkill /f /im python.exe
#taskkill /f /im node.exe
#docker compose up -d --build
#docker compose down && docker compose up -d --build
#http://127.0.0.1:8080/
#docker compose logs -f backend
#ssh root@47.86.108.159
#http://47.86.108.159:8080

import json
import os
import sys
import uuid
import asyncio
import re
import traceback
from contextlib import asynccontextmanager
from concurrent.futures import ThreadPoolExecutor

import io
import tempfile
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

from src.schemas.resume import (
    RadarMetrics,
    OptimizeMode,
    ResumeOptimizeRequest,
    ChatRequest,
)
from src.graph import build_graph
from src.routes.agent_router import router as agent_router
from src.auth.router import router as auth_router
from src.nodes.visual_payload import compile_to_visual_payload
from langgraph.checkpoint.sqlite.aio import AsyncSqliteSaver

# ── 启动时加载 .env ──
from dotenv import load_dotenv
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__) or ".", ".env"))

# ── v5.0 持久化状态机：_app_graph 在 lifespan 中用 with SqliteSaver 解包后延迟编译 ──
_app_graph = None  # lifespan 中通过 build_graph(checkpointer) 完成注入

# ── v5.0 全局单例线程池：根除每请求建池的线程风暴 ──
GLOBAL_POOL = ThreadPoolExecutor(max_workers=20)

# ═══════════════════════════════════════════════════════════════
# v5.2 三道安全防线：限流器(含IP白名单) + JWT Bearer Token + CORS 白名单
# ═══════════════════════════════════════════════════════════════

# ── 防线〇：Slowapi 限流豁免白名单（Docker网桥 + 本地回环免限流误杀）──
EXEMPT_IPS = {"127.0.0.1", "localhost", "172.19.0.1"}

def _rate_limit_key(request: Request):
    """自定义限流键函数：白名单 IP 返回 None（豁免），其余走默认 IP 提取。

    slowapi 内部对 key_func 返回 None 的请求完全跳过速率检查，
    防止容器内网自调、本地调试等场景触发限流误杀。
    """
    client_ip = get_remote_address(request)
    if client_ip in EXEMPT_IPS:
        return None
    return client_ip

limiter = Limiter(key_func=_rate_limit_key)

# ── 防线一：JWT Bearer Token 全局鉴权（v5.2 正规军认证系统）──
from src.auth.security import get_current_user


# ── SSE 工具函数 ──

def _sse_event(event: str, data: dict) -> str:
    """构建一条 SSE (Server-Sent Event) 消息帧"""
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _build_radar(dims: dict, declared_score: int) -> "RadarMetrics | None":
    """从维度分数字典构建 RadarMetrics，含分项一致性校验"""
    if not dims:
        return None
    # ── v5.9 None 安全兜底：防止上游节点传入 None 值导致 TypeError ──
    _safe_score = declared_score if declared_score is not None else 0
    dims_sum = dims.get("jd_match", 0) + dims.get("star_completion", 0) + dims.get("verb_quality", 0)
    if _safe_score > 0 and abs(_safe_score - dims_sum) <= 3:
        total = _safe_score
    elif dims_sum > 0:
        total = dims_sum
    else:
        total = _safe_score
    return RadarMetrics.from_dimensions(dims, total)


def _build_stress_questions(raw_questions: list) -> list[dict]:
    """将 interviewer 节点返回的原始压测题转为前端可用的字典列表"""
    result = []
    for q in raw_questions:
        try:
            result.append({
                "question_number": q.get("question_number", len(result) + 1),
                "category": q.get("category", "技术深度"),
                "question": q.get("question", ""),
                "expected_points": q.get("expected_points", []),
            })
        except Exception as e:
            print(f"[api] 压测题解析失败: {e}")
    return result


def _safe_list(dims: dict, key: str) -> list[str]:
    """从 dimension_scores 字典安全提取列表字段并转字符串列表"""
    val = dims.get(key, [])
    if isinstance(val, list):
        return [str(v) for v in val]
    if isinstance(val, str):
        return [val] if val else []
    return []


# ── FastAPI 应用工厂 ──

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期：AsyncSqliteSaver 异步解包 → 双图编译 → ChromaDB 预热"""
    global _app_graph

    print("[server] AI-Resume-Evolver v5.2 启动中...")

    # ── v5.2 AsyncSqliteSaver 异步上下文管理器解包，生命周期 = 进程存活期 ──
    os.makedirs("data", exist_ok=True)
    async with AsyncSqliteSaver.from_conn_string("data/resume_state.db") as checkpointer:
        # ── v7.1 信任 AsyncSqliteSaver 内部的 SQLite 连接管理；
        # WAL mode 是持久属性（设置后写回 db 文件头），无需额外 PRAGMA 连接 ──
        try:
            import sqlite3
            _c = sqlite3.connect("file:data/resume_state.db?mode=ro", uri=True, timeout=2)
            _mode = _c.execute("PRAGMA journal_mode;").fetchone()
            _c.close()
            print(f"[server] SQLite journal_mode={_mode[0] if _mode else 'unknown'}")
        except Exception:
            print("[server] SQLite journal_mode 探测跳过（AsyncSqliteSaver 内部已接管）")

        _app_graph = build_graph(checkpointer=checkpointer)
        print("[server] AsyncSqliteSaver 异步状态机已挂载 → data/resume_state.db")

        # ── v5.2 认证种子守卫：自动建表 + admin 账号注入 ──
        try:
            from src.auth.seed import ensure_users_table_and_admin
            ensure_users_table_and_admin()
        except Exception as e:
            print(f"[server] 认证种子注入失败（非致命）: {e}")

        print("[server] 预热 ChromaDB 连接 + 种子数据守卫...")
        try:
            from src.config import get_vector_db_client, get_collection_name, ensure_seed_data
            seeded = ensure_seed_data()
            if seeded > 0:
                print(f"[server] 种子数据守卫已激活：自动灌入 {seeded} 条金牌案例")
            client = get_vector_db_client()
            collection = client.get_or_create_collection(name=get_collection_name())
            count = collection.count()
            print(f"[server] ChromaDB 就绪，Collection '{get_collection_name()}' 共 {count} 条向量")
        except Exception as e:
            print(f"[server] ChromaDB 预热失败（非致命）: {e}")

        # ── v5.0 双图共用同一个 SqliteSaver，agent graph 延迟注入 ──
        from src.graphs.agent_graph import init_agent_checkpointer
        init_agent_checkpointer(checkpointer)
        print("[server] Agent Graph checkpointer 注入完成，双图共用 data/resume_state.db")

        print("[server] 服务已就绪，接受请求。")
        yield
        print("[server] 服务关闭。")


app = FastAPI(
    title="AI-Resume-Evolver API",
    description="AI 简历智能优化引擎 —— 一键生成 + 多智能体博弈 + MockInterviewer 压测 + 多轮交互",
    version="5.1.0",
    lifespan=lifespan,
    dependencies=[Depends(get_current_user)],  # v5.2 JWT Bearer Token 全局鉴权
)

# ── v5.1 CORS 白名单（从环境变量读取，杜绝 allow_origins=["*"]）──
_allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "http://localhost:8080,http://127.0.0.1:8080")
ALLOWED_ORIGINS = [origin.strip() for origin in _allowed_origins_str.split(",") if origin.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── v5.1 Slowapi 限流器注册 + 429 异常处理器 ──
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, lambda request, exc: JSONResponse(
    status_code=429,
    content={"detail": "[限流熔断] 请求过于频繁，请等待 1 分钟后再试。"},
))


# ── 健康检查 ──

@app.get("/health")
@limiter.limit("60/minute")
async def health_check(request: Request):
    return {"status": "ok", "version": "5.1.0", "service": "AI-Resume-Evolver"}


# ── SSE 流式管道 ──

async def _stream_pipeline(initial_state: dict, thread_id: str = ""):
    """
    SSE 事件流生成器 —— 在 ThreadPoolExecutor 中运行 LangGraph 全链路，
    通过 asyncio.Queue 桥接同步图执行与异步 SSE 推送。

    使用应用级共享图 _app_graph，通过 thread_id 启用 MemorySaver 断点续传。
    一键模式初始 user_supplement 为空，entry_router 自动走 retriever 全链路。

    事件序列:
      1. radar_init:   初筛完成 → 原始简历 6-3-1 雷达指标
      2. resume_stream: Editor 完成 → 精修简历全文
      3. final:         Evaluator + Interviewer 完成 → 终评雷达 + 压测题
      4. done:          流结束
      5. error:         异常（如有）
    """
    loop = asyncio.get_running_loop()
    queue: asyncio.Queue = asyncio.Queue()
    config = {"configurable": {"thread_id": thread_id}} if thread_id else {}

    def _run_graph():
        """在子线程中同步执行 LangGraph stream，将每次状态快照推入队列"""
        try:
            for state in _app_graph.stream(initial_state, config=config, stream_mode="values"):
                loop.call_soon_threadsafe(queue.put_nowait, ("state", state))
            loop.call_soon_threadsafe(queue.put_nowait, ("done", None))
        except Exception as e:
            loop.call_soon_threadsafe(queue.put_nowait, ("error", e))

    yielded_radar = False
    yielded_resume = False
    final_state = None

    GLOBAL_POOL.submit(_run_graph)

    while True:
        status, value = await queue.get()

        if status == "error":
            traceback.print_exc()
            yield _sse_event("error", {
                "error": f"{type(value).__name__}: {str(value)[:200]}"
            })
            yield _sse_event("done", {})
            return

        if status == "done":
            break

        # status == "state": 处理每次节点完成后的全量状态快照
        state = value
        final_state = state

        # ── 里程碑 1: PreEvaluator 完成 → 推送原始简历雷达 + 诊断原文 ──
        if not yielded_radar and state.get("pre_eval_dimensions"):
            yielded_radar = True
            dims = state["pre_eval_dimensions"]
            pre_total = dims.get("jd_match", 0) + dims.get("star_completion", 0) + dims.get("verb_quality", 0)
            if pre_total == 0:
                pre_total = state.get("score", 0)
            radar = RadarMetrics.from_dimensions(dims, pre_total)

            # ── v4.6 诊断原文：从 pre_eval_dimensions 提取结构化诊断数据 ──
            diagnosis = {
                "feedback": state.get("evaluation_feedback", ""),
                "core_tool_overlap": state.get("node_status", ""),
                "matched_skills": _safe_list(dims, "matched_skills"),
                "missing_skills": _safe_list(dims, "missing_skills"),
                "star_strengths": _safe_list(dims, "star_strengths"),
                "star_weaknesses": _safe_list(dims, "star_weaknesses"),
                "weak_verbs": _safe_list(dims, "weak_verbs"),
            }

            yield _sse_event("radar_init", {
                "original_resume_radar": {
                    "jd_matching_score": radar.jd_matching_score,
                    "star_perf_score": radar.star_perf_score,
                    "action_verbs_score": radar.action_verbs_score,
                    "total_score": radar.total_score,
                },
                "diagnosis": diagnosis,
            })
            print(f"[sse] 分帧1 radar_init: 原始雷达 {pre_total}/100 "
                  f"(JD: {radar.jd_matching_score}/60, STAR: {radar.star_perf_score}/30, "
                  f"Verb: {radar.action_verbs_score}/10) | "
                  f"覆盖技能: {len(diagnosis['matched_skills'])}项, "
                  f"缺失: {len(diagnosis['missing_skills'])}项")

        # ── 里程碑 2: Editor 完成 → 推送精修简历文本 + 混合解耦载荷 ──
        if not yielded_resume and state.get("revised_resume"):
            yielded_resume = True
            from src.utils.text_sanitizer import sanitize_resume_text
            revised = sanitize_resume_text(state["revised_resume"], log_prefix="[sse]")
            opt_summary = state.get("optimization_summary", "")
            clean_json = state.get("clean_resume_json", {})
            vp = compile_to_visual_payload(revised)
            yield _sse_event("resume_stream", {
                "optimized_resume_text": revised,
                "text_length": len(revised),
                "optimization_summary": opt_summary,
                "clean_resume_json": clean_json,
                "visual_payload": vp,
            })
            print(f"[sse] 分帧2 resume_stream: 精修文本 {len(revised)} 字符, "
                  f"skills={len(vp.get('skills', []))}项, "
                  f"summary {len(opt_summary)} 字符, json_keys {list(clean_json.keys()) if clean_json else '[]'}")

    # ── 里程碑 3: 全链路完成 → 推送终评雷达 + 压测题 ──
    if final_state:
        eval_dims = final_state.get("eval_dimensions", {})
        eval_score = final_state.get("score", 0)
        opt_radar = _build_radar(eval_dims, eval_score)

        questions = _build_stress_questions(
            final_state.get("stress_test_questions", [])
        )

        pre_eval_dims = final_state.get("pre_eval_dimensions", {})
        pre_total = (pre_eval_dims.get("jd_match", 0) +
                     pre_eval_dims.get("star_completion", 0) +
                     pre_eval_dims.get("verb_quality", 0))

        is_extreme_gap = final_state.get("difficulty_flag", "") == "EXTREME_GAP"
        from src.utils.text_sanitizer import sanitize_resume_text
        revised_final = sanitize_resume_text(final_state.get("revised_resume", ""), log_prefix="[sse-final]")
        vp_final = compile_to_visual_payload(revised_final) if revised_final else {}

        # ── v5.4 极端情况熔断器：防止用户信任危机 ──
        # v5.9 None 安全兜底：上游节点可能传入 None score
        _safe_eval_score = eval_score if eval_score is not None else 0
        _safe_pre_total = pre_total if pre_total is not None else 0
        score_improvement = _safe_eval_score - _safe_pre_total if _safe_pre_total > 0 else 0
        display_score_change = True
        circuit_breaker_triggered = False

        _CIRCUIT_BREAKER_MSG = (
            "当前简历与目标 JD 存在明显的技术栈脱节，"
            "系统已尽力优化表达，但核心硬实力差距较大，"
            "建议针对性提升相关技术、多积累项目经验后再尝试对齐。"
        )

        if (score_improvement < 5) and _safe_eval_score < 50:
            circuit_breaker_triggered = True
            display_score_change = False
            score_improvement = 0
            print(f"[sse] ⚠️ 信任熔断触发 | 原始={_safe_pre_total} 优化后={_safe_eval_score} "
                  f"delta={_safe_eval_score - _safe_pre_total if _safe_pre_total > 0 else 'N/A'} | "
                  f"已抑制分数变化展示，重写评语为专业建议")

        yield _sse_event("final", {
            "optimized_resume_radar": {
                "jd_matching_score": opt_radar.jd_matching_score if opt_radar else 0,
                "star_perf_score": opt_radar.star_perf_score if opt_radar else 0,
                "action_verbs_score": opt_radar.action_verbs_score if opt_radar else 0,
                "total_score": opt_radar.total_score if opt_radar else eval_score,
            },
            "optimized_resume_text": revised_final,
            "stress_test_questions": questions,
            "difficulty_flag": final_state.get("difficulty_flag", ""),
            "is_extreme_gap": is_extreme_gap,
            "iteration_count": final_state.get("iteration_count", 0),
            "score_improvement": score_improvement,
            "display_score_change": display_score_change,
            "circuit_breaker_triggered": circuit_breaker_triggered,
            "internal_monologue": final_state.get("internal_monologue", ""),
            "evaluation_feedback": (
                _CIRCUIT_BREAKER_MSG if circuit_breaker_triggered
                else final_state.get("evaluation_feedback", "")
            ),
            "pre_eval_dimensions": pre_eval_dims,
            "eval_dimensions": eval_dims,
            "optimization_summary": (
                _CIRCUIT_BREAKER_MSG if circuit_breaker_triggered
                else final_state.get("optimization_summary", "")
            ),
            "clean_resume_json": final_state.get("clean_resume_json", {}),
            "visual_payload": vp_final,
            "session_id": thread_id,
        })
        print(f"[sse] 分帧3 final: 终评 {eval_score}/100, 压测题 {len(questions)} 道, "
              f"提升 +{score_improvement} 分"
              f"{' [信任熔断已激活]' if circuit_breaker_triggered else ''}")

    yield _sse_event("done", {})
    print(f"[sse] 流式推送完成 (4/4 帧已全部发送), thread_id={thread_id}")


# ═══════════════════════════════════════════════════════════════
# v5.5 输入保护性熔断：极短/无效输入优雅降级
# ═══════════════════════════════════════════════════════════════

_INPUT_DEGRADED_MSG = (
    "当前输入的简历或 JD 内容过于简略，技术栈严重脱节。"
    "系统已启动保护性熔断，请补充具体的项目经历与核心技术栈后再尝试演进。"
)

_MIN_RESUME_CHARS = 10  # 去空格后最少字符数
_MIN_JD_CHARS = 10


async def _generate_degraded_sse(thread_id: str, reason: str) -> str:
    """当输入文本过短时，跳过全链路直接返回保护性熔断 SSE 帧。

    前端收到 circuit_breaker_triggered=true 后展示琥珀色提示卡。
    """
    yield _sse_event("radar_init", {
        "original_resume_radar": {
            "jd_matching_score": 0,
            "star_perf_score": 0,
            "action_verbs_score": 0,
            "total_score": 0,
        },
        "diagnosis": {
            "feedback": reason,
            "core_tool_overlap": "输入保护性熔断",
            "matched_skills": [],
            "missing_skills": [],
            "star_strengths": [],
            "star_weaknesses": [],
            "weak_verbs": [],
        },
    })

    yield _sse_event("final", {
        "optimized_resume_radar": {
            "jd_matching_score": 0,
            "star_perf_score": 0,
            "action_verbs_score": 0,
            "total_score": 0,
        },
        "optimized_resume_text": "",
        "stress_test_questions": [],
        "difficulty_flag": "DEGRADED",
        "is_extreme_gap": False,
        "iteration_count": 0,
        "score_improvement": 0,
        "display_score_change": False,
        "circuit_breaker_triggered": True,
        "internal_monologue": "",
        "evaluation_feedback": reason,
        "pre_eval_dimensions": {},
        "eval_dimensions": {},
        "optimization_summary": reason,
        "clean_resume_json": {},
        "visual_payload": None,
        "session_id": thread_id,
    })

    yield _sse_event("done", {})
    print(f"[sse] 输入保护性熔断完成, thread_id={thread_id}")


# ── 核心路由：简历优化 ──

@app.post("/api/v1/resume/optimize")
@limiter.limit("5/minute")
async def optimize_resume(request: Request, payload: ResumeOptimizeRequest):
    """
    简历优化接口 —— SSE 流式事件推送。

    **ONE_CLICK 模式** (StreamingResponse, text/event-stream):
    依次推送 4 帧 SSE 事件，前端可实时渲染阶段性进度：

    | 事件名 | 触发时机 | data 负载 |
    |--------|----------|-----------|
    | `radar_init` | PreEvaluator 初筛完成 | `original_resume_radar`: 原始简历 6-3-1 雷达指标，前端渲染初始雷达图 |
    | `resume_stream` | Editor 精修完成 | `optimized_resume_text`: 精修后完整简历文本，前端打字机渐进显示 |
    | `final` | Evaluator+Interviewer 完成 | `optimized_resume_radar`: 终评雷达, `stress_test_questions`: 3 道面试压测题, `score_improvement`: 提升分, `internal_monologue`: 毒舌批评, `session_id`: 会话 ID 供后续 chat 续接 |
    | `done` | 流正常结束 | 空 JSON `{}`，前端关闭 EventSource |
    | `error` | 异常中断 | `error`: 异常信息，随后推送 `done` 关闭流 |

    **INTERACTIVE 模式**: 已并网！前端应在拿到 session_id 后，通过 /api/v1/resume/chat 发起多轮补充。
    此端点返回的 final 帧中包含 session_id，前端需保存该 ID 用于后续 chat 调用。
    """
    # ── v5.6 JWT 强制绑定：user_id 从令牌推导，防御跨用户身份伪造 ──
    auth_user = getattr(request.state, "user", None)
    user_id = auth_user.username if auth_user else payload.user_id
    if auth_user and payload.user_id and payload.user_id != auth_user.username:
        print(f"[Security] ⚠️ 令牌不匹配拦截: payload.user_id={payload.user_id} "
              f"vs JWT={auth_user.username} — 已强制使用 JWT 身份")
    # ── v7.1 会话级 thread_id: 每次一键优化生成全新 UUID，彻底隔离不同批次的记忆 ──
    resume_id = payload.resume_id if payload.resume_id and payload.resume_id != "default_resume" else f"res_{uuid.uuid4().hex[:8]}"
    thread_id = f"pipeline::{user_id}::{uuid.uuid4().hex[:12]}"
    print(f"[记忆沙箱点火] 全新会话已锁定: {thread_id} (resume_id={resume_id})")

    # ── v5.5 输入保护性熔断：极短/无效输入优雅降级，杜绝 422 ──
    stripped_resume = payload.resume_text.strip()
    stripped_jd = payload.jd_text.strip()
    if len(stripped_resume) < _MIN_RESUME_CHARS or len(stripped_jd) < _MIN_JD_CHARS:
        short_parts = []
        if len(stripped_resume) < _MIN_RESUME_CHARS:
            short_parts.append(f"简历仅 {len(stripped_resume)} 字符（阈值 {_MIN_RESUME_CHARS}）")
        if len(stripped_jd) < _MIN_JD_CHARS:
            short_parts.append(f"JD 仅 {len(stripped_jd)} 字符（阈值 {_MIN_JD_CHARS}）")
        print(f"[api] ⚠️ 输入保护性熔断触发 | {', '.join(short_parts)} | thread_id={thread_id}")

        return StreamingResponse(
            _generate_degraded_sse(thread_id, _INPUT_DEGRADED_MSG),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )

    print(f"\n[api] 收到优化请求 [{payload.mode.value}]: 简历 {len(payload.resume_text)} 字符, "
          f"JD {len(payload.jd_text)} 字符, thread_id={thread_id}")

    initial_state = {
        "user_id": user_id,
        "resume_id": resume_id,
        "resume": payload.resume_text,
        "jd": payload.jd_text,
        "rag_context": "",
        "revised_resume": "",
        "internal_monologue": "",
        "tool_outputs": [],
        "score": 0,
        "evaluation_feedback": "",
        "iteration_count": 0,
        "difficulty_flag": "",
        "node_status": "",
        "pre_eval_dimensions": {},
        "eval_dimensions": {},
        "stress_test_questions": [],
        "optimization_summary": "",
        "clean_resume_json": {},
        "chat_history": [],
        "user_supplement": "",
        "session_id": "",      # 一键模式不注入 session_id，eval_condition 走正常分诊
        "turn_count": 0,
        "step_count": 0,
        "total_tokens": 0,
    }

    return StreamingResponse(
        _stream_pipeline(initial_state, thread_id=thread_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── 交互模式 SSE 流式管道 ──

async def _stream_chat_pipeline(thread_id: str, user_message: str):
    """
    交互模式 SSE 事件流生成器 —— 接收用户补充信息，通过共享图 _app_graph
    在已有 session 的 checkpoint 上续跑，走 chat_editor → evaluator 闭环。

    事件序列:
      1. status:       chat_editor 完成 → node_status 状态帧
      2. resume_stream: 更新后的简历全文 (含 node_status)
      3. final:         Evaluator 评分 + Interviewer 压测题
      4. done:          流结束
      5. error:         异常（如有）
    """
    loop = asyncio.get_running_loop()
    queue: asyncio.Queue = asyncio.Queue()
    config = {"configurable": {"thread_id": thread_id}}

    # ── v7.1 user_id / resume_id 不再从 thread_id 解析，改为从 checkpoint 继承 ──
    initial_input = {
        "user_supplement": user_message,
        "session_id": thread_id,  # 注入 session_id 以激活交互模式路由
        "step_count": 0,
        "total_tokens": 0,
    }

    print(f"[记忆沙箱点火] 线程已锁定: {thread_id}")
    print(f"[chat_sse] 启动交互流: thread_id={thread_id}, "
          f"user_message={len(user_message)} 字符")

    def _run_graph():
        try:
            for state in _app_graph.stream(initial_input, config=config, stream_mode="values"):
                loop.call_soon_threadsafe(queue.put_nowait, ("state", state))
            loop.call_soon_threadsafe(queue.put_nowait, ("done", None))
        except Exception as e:
            loop.call_soon_threadsafe(queue.put_nowait, ("error", e))

    final_state = None
    last_resume_len = 0

    GLOBAL_POOL.submit(_run_graph)

    while True:
        status, value = await queue.get()

        if status == "error":
            traceback.print_exc()
            yield _sse_event("error", {
                "error": f"{type(value).__name__}: {str(value)[:200]}"
            })
            yield _sse_event("done", {})
            return

        if status == "done":
            break

        state = value
        final_state = state

        # ── 帧: node_status 状态推送 (对齐前端 status case) ──
        node_status = state.get("node_status", "")
        if node_status:
            yield _sse_event("status", {
                "node_status": node_status,
                "turn_count": state.get("turn_count", 0),
            })
            print(f"[chat_sse] status: {node_status[:80]}...")

        # ── 帧: 简历文本变更推送 (对齐前端 resume_stream case) ──
        from src.utils.text_sanitizer import sanitize_resume_text
        revised = sanitize_resume_text(state.get("revised_resume", ""), log_prefix="[chat_sse]")
        if revised and len(revised) != last_resume_len:
            last_resume_len = len(revised)
            vp = compile_to_visual_payload(revised)
            yield _sse_event("resume_stream", {
                "optimized_resume_text": revised,
                "text_length": len(revised),
                "turn_count": state.get("turn_count", 0),
                "node_status": node_status or "",
                "visual_payload": vp,
            })
            print(f"[chat_sse] resume_stream: {len(revised)} 字符, "
                  f"skills={len(vp.get('skills', []))}项")

    # ── 终帧: 评分 + 压测题 + 混合解耦载荷 ──
    if final_state:
        eval_dims = final_state.get("eval_dimensions", {})
        eval_score = final_state.get("score", 0)
        opt_radar = _build_radar(eval_dims, eval_score)

        questions = _build_stress_questions(
            final_state.get("stress_test_questions", [])
        )

        revised_final = sanitize_resume_text(final_state.get("revised_resume", ""), log_prefix="[chat_sse-final]")
        vp_final = compile_to_visual_payload(revised_final) if revised_final else {}

        yield _sse_event("final", {
            "optimized_resume_radar": {
                "jd_matching_score": opt_radar.jd_matching_score if opt_radar else 0,
                "star_perf_score": opt_radar.star_perf_score if opt_radar else 0,
                "action_verbs_score": opt_radar.action_verbs_score if opt_radar else 0,
                "total_score": opt_radar.total_score if opt_radar else eval_score,
            },
            "optimized_resume_text": revised_final,
            "stress_test_questions": questions,
            "difficulty_flag": final_state.get("difficulty_flag", ""),
            "turn_count": final_state.get("turn_count", 0),
            "internal_monologue": final_state.get("internal_monologue", ""),
            "visual_payload": vp_final,
            "session_id": thread_id,
        })
        print(f"[chat_sse] final: 终评 {eval_score}/100, "
              f"skills={len(vp_final.get('skills', []))}项, "
              f"轮次 {final_state.get('turn_count', 0)}, 压测题 {len(questions)} 道")

    yield _sse_event("done", {})
    print(f"[chat_sse] 交互流推送完成, thread_id={thread_id}")


# ── 交互模式路由：多轮对话 ──

@app.post("/api/v1/resume/chat")
@limiter.limit("10/minute")
async def chat_resume(request: Request, payload: ChatRequest):
    """
    多轮对话交互编辑接口 —— SSE 流式事件推送。

    前端需先通过 /api/v1/resume/optimize 获取 session_id，
    再将该 session_id 作为 thread_id 传入本端点进行多轮补充。

    **事件帧序列**:

    | 事件名 | 触发时机 | data 负载 |
    |--------|----------|-----------|
    | `status` | 节点状态变更 | `node_status`: 当前阶段描述, `turn_count`: 当前轮次 |
    | `resume_stream` | chat_editor 完成增量编辑 | `optimized_resume_text`: 更新后简历全文, `text_length`: 字符数, `node_status`: 状态信息 |
    | `final` | Evaluator+Interviewer 完成 | `optimized_resume_radar`: 终评雷达, `stress_test_questions`: 压测题, `turn_count`: 总轮次 |
    | `done` | 流正常结束 | 空 JSON `{}` |
    | `error` | 异常中断 | `error`: 异常信息 |
    """
    # ── v7.1 JWT 强制绑定：从 checkpoint 读取会话归属，而非从 thread_id 字符串解析 ──
    auth_user = getattr(request.state, "user", None)
    if auth_user:
        try:
            config = {"configurable": {"thread_id": payload.thread_id}}
            snapshot = await _app_graph.aget_state(config)
            if snapshot and snapshot.values:
                checkpoint_owner = (snapshot.values.get("user_id") or "").strip()
                if checkpoint_owner and checkpoint_owner != auth_user.username:
                    raise HTTPException(
                        status_code=status.HTTP_403_FORBIDDEN,
                        detail=f"[安全熔断] 无权访问他人会话: checkpoint 归属 {checkpoint_owner}, JWT 身份 {auth_user.username}",
                    )
        except HTTPException:
            raise
        except Exception:
            pass  # checkpoint 不存在时放行，图内部会做空状态优雅降级

    print(f"\n[api] 收到交互补充请求: thread_id={payload.thread_id}, "
          f"message={len(payload.user_message)} 字符")

    return StreamingResponse(
        _stream_chat_pipeline(payload.thread_id, payload.user_message),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )

# ── 多模态文件上传解析端点 ──

@app.post("/api/v1/upload/parse")
@limiter.limit("10/minute")
async def upload_and_parse(request: Request, file: UploadFile = File(...)):
    """
    文件/截屏上传解析端点 —— 支持 PDF/DOCX/TXT 文件物理去噪 + 图片多模态 Vision OCR。

    **支持的文件类型**:
    - `.pdf`  → pypdf 逐页提取文本
    - `.docx` → docx2txt 结构化文本提取
    - `.txt` / `.md` → 直接 UTF-8 读取
    - `.png` / `.jpg` / `.jpeg` / `.webp` / `.gif` / `.bmp` → DeepSeek 多模态视觉解析

    **响应体**:
    ```json
    {"success": true, "text": "提取的纯文本", "file_type": "pdf|docx|txt|image", "filename": "原始文件名"}
    ```
    """
    filename = file.filename or "unknown"
    ext = Path(filename).suffix.lower()

    # 文件大小安全校验：上限 20MB
    MAX_SIZE = 20 * 1024 * 1024
    raw_bytes = await file.read()
    if len(raw_bytes) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="文件过大，请上传 20MB 以内的文件")
    if len(raw_bytes) == 0:
        raise HTTPException(status_code=400, detail="文件为空，请重新选择文件")

    print(f"[upload] 收到上传文件: {filename} ({len(raw_bytes)} bytes, type={ext})")

    try:
        # ── 纯文本类 ──
        if ext in (".txt", ".md"):
            text = raw_bytes.decode("utf-8").strip()
            print(f"[upload] TXT 解析完成: {len(text)} 字符")
            return {"success": True, "text": text, "file_type": "txt", "filename": filename}

        # ── PDF ──
        if ext == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            pages: list[str] = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages).strip()
            if not text:
                raise HTTPException(status_code=422, detail="PDF 无法提取文本，可能是扫描件或图片型 PDF，请尝试截图后用图片上传")
            print(f"[upload] PDF 解析完成: {len(text)} 字符, {len(pages)} 页")
            return {"success": True, "text": text, "file_type": "pdf", "filename": filename}

        # ── DOCX ──
        if ext == ".docx":
            import docx2txt

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(raw_bytes)
                tmp_path = tmp.name
            try:
                text = docx2txt.process(tmp_path)
                text = (text or "").strip()
            finally:
                os.unlink(tmp_path)
            if not text:
                raise HTTPException(status_code=422, detail="DOCX 文件无法提取文本，请检查文件内容")
            print(f"[upload] DOCX 解析完成: {len(text)} 字符")
            return {"success": True, "text": text, "file_type": "docx", "filename": filename}

        # ── 图片 / 截屏（v4.4：DeepSeek 多模态视觉 Context 注入，彻底淘汰 EasyOCR）──
        if ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"):
            from src.utils.visual_processor import parse_resume_image_via_vlm

            print(f"[upload] 启动 DeepSeek 多模态视觉解析: {filename} ({len(raw_bytes)} bytes)")
            text = await parse_resume_image_via_vlm(raw_bytes)
            text = (text or "").strip()
            if not text:
                raise HTTPException(status_code=422, detail="图片中未识别到文字内容，请确认图片包含清晰的简历或 JD 文字")
            print(f"[upload] DeepSeek 视觉解析完成: {len(text)} 字符")
            return {"success": True, "text": text, "file_type": "image", "filename": filename}

        # ── 不支持的类型 ──
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型 ({ext})，请上传 PDF、DOCX、TXT 或图片（PNG/JPG/WEBP/BMP）",
        )

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)[:200]}")


# ═══════════════════════════════════════════════════════════════
# v5.3 DOCX 导出：POST /api/resume/export/docx
# ═══════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Pt, Inches

@app.post("/api/resume/export/docx")
async def export_resume_docx(request: Request):
    """
    将 Markdown 简历内容导出为 A4 规格 DOCX 文件。
    仅对 **bold** 文本加粗，不添加任何边框/背景/色块。
    """
    body = await request.json()
    markdown_content = body.get("markdown_content", "")

    if not markdown_content or not markdown_content.strip():
        raise HTTPException(status_code=400, detail="markdown_content 不能为空")

    doc = Document()

    # A4 标准边距：上下左右各 1 英寸
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = markdown_content.split("\n")

    for line in lines:
        line = line.rstrip()
        if not line:
            continue

        # ### 小标题 → 13pt Arial Bold
        if line.startswith("### "):
            text = line[4:].strip()
            text = re.sub(r"\*\*", "", text)  # 清洗行内 ** 标记
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(13)
            run.font.name = "Arial"
            run.bold = True
            continue

        # ## 中标题 → 15pt Arial Bold
        if line.startswith("## "):
            text = line[3:].strip()
            text = re.sub(r"\*\*", "", text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(15)
            run.font.name = "Arial"
            run.bold = True
            continue

        # # 大标题 → 18pt Arial Bold
        if line.startswith("# "):
            text = line[2:].strip()
            text = re.sub(r"\*\*", "", text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.font.name = "Arial"
            run.bold = True
            continue

        # - 或 * 列表项
        if re.match(r"^[\-\*]\s+", line):
            list_text = re.sub(r"^[\-\*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_docx_inline_runs(p, list_text)
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_docx_inline_runs(p, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=resume.docx"},
    )


def _add_docx_inline_runs(paragraph, text: str):
    """
    解析 **bold** 标记，仅加粗。
    严禁添加方框、边框、背景底色、胶囊气泡或任何装饰。
    """
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ── v5.2 JWT 认证路由：POST /api/v1/auth/login ──
app.include_router(auth_router)

# ── Agent 模式路由：LangGraph ReAct 中央大脑 SSE 流式端点 ──
app.include_router(agent_router)

if __name__ == "__main__":
    import uvicorn
    print("启动 AI-Resume-Evolver FastAPI 服务 v5.0 (双模拓扑 + SqliteSaver + 全局线程池)...")
    uvicorn.run(
        "main:app",
        host=os.getenv("HOST", "127.0.0.1"),
        port=int(os.getenv("PORT", "8001")),
        reload=False,
        log_level="info",
        timeout_keep_alive=300,  # 长连接保活 5 分钟（覆盖串行 LLM 调用的完整链路）
        timeout_graceful_shutdown=30,
    )
