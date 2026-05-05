import os
import re
import textwrap
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

_FONT_DIRS = [
    os.path.join(os.path.dirname(__file__), "..", "..", "assets", "fonts"),
    "C:/Windows/Fonts",
]

_FONT_CANDIDATES = {
    "msyh": ["msyh.ttc", "msyh.ttf"],
    "msyhbd": ["msyhbd.ttc", "msyhbd.ttf"],
    "simsun": ["simsun.ttc", "simsun.ttf"],
    "simhei": ["simhei.ttf"],
}


def _find_font(name: str) -> str | None:
    candidates = _FONT_CANDIDATES.get(name, [f"{name}.ttf", f"{name}.ttc"])
    for base_dir in _FONT_DIRS:
        if not os.path.isdir(base_dir):
            continue
        for fname in candidates:
            path = os.path.join(base_dir, fname)
            if os.path.isfile(path):
                return path
    return None


def _register_cn_fonts():
    if "CNFont" in pdfmetrics._fonts:
        return

    font_path = _find_font("msyh") or _find_font("simsun") or _find_font("simhei")
    if font_path is None:
        print("[exporter] 警告：未找到中文字体文件，PDF 可能无法正常显示中文")
        print("[exporter] 请将 simsun.ttf 或 msyh.ttc 放入 assets/fonts/ 目录")
        return

    try:
        pdfmetrics.registerFont(TTFont("CNFont", font_path))
        pdfmetrics.registerFont(TTFont("CNFontBold", _find_font("msyhbd") or font_path))
        print(f"[exporter] 中文字体已注册: {os.path.basename(font_path)}")
    except Exception as e:
        print(f"[exporter] 字体注册失败: {e}")


_CN_FONT_REGISTERED = False


def _ensure_cn_font():
    global _CN_FONT_REGISTERED
    if not _CN_FONT_REGISTERED:
        _register_cn_fonts()
        _CN_FONT_REGISTERED = True


_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
_MD_UNORDERED = re.compile(r"^[-*+]\s+(.+)$")
_MD_ORDERED = re.compile(r"^\d+[.)]\s+(.+)$")


class _MarkdownBlock:
    def __init__(self, block_type: str, text: str = "", items: list[str] | None = None):
        self.block_type = block_type
        self.text = text
        self.items = items or []


def _parse_markdown(md_content: str) -> list[_MarkdownBlock]:
    lines = md_content.strip().split("\n")
    blocks: list[_MarkdownBlock] = []
    pending_list: list[str] = []
    list_type: str | None = None

    def _flush_list():
        nonlocal pending_list, list_type
        if pending_list:
            blocks.append(_MarkdownBlock(list_type or "ul", items=list(pending_list)))
            pending_list = []
            list_type = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            _flush_list()
            continue

        heading_match = _MD_HEADING.match(stripped)
        if heading_match:
            _flush_list()
            level = len(heading_match.group(1))
            blocks.append(_MarkdownBlock(f"h{level}", text=heading_match.group(2)))
            continue

        ul_match = _MD_UNORDERED.match(stripped)
        if ul_match:
            if list_type and list_type != "ul":
                _flush_list()
            list_type = "ul"
            pending_list.append(ul_match.group(1))
            continue

        ol_match = _MD_ORDERED.match(stripped)
        if ol_match:
            if list_type and list_type != "ol":
                _flush_list()
            list_type = "ol"
            pending_list.append(ol_match.group(1))
            continue

        _flush_list()
        blocks.append(_MarkdownBlock("p", text=stripped))

    _flush_list()
    return blocks


def _split_bold_segments(text: str) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    pos = 0
    for match in _MD_BOLD.finditer(text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False))
        segments.append((match.group(1), True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False))
    if not segments:
        segments.append((text, False))
    return segments


class ResumeExporter:
    def __init__(self):
        self._cn_font_ready = False

    def _init_cn_font(self):
        if not self._cn_font_ready:
            _ensure_cn_font()
            self._cn_font_ready = True

    def to_docx(self, md_content: str, output_path: str):
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        pf = style.paragraph_format
        pf.line_spacing = 1.25
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        for h_level in range(1, 5):
            hs = doc.styles[f"Heading {h_level}"]
            hs.font.name = "微软雅黑"
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            hs.font.color.rgb = RGBColor(0, 0, 0)
            if h_level == 1:
                hs.font.size = Pt(16)
            elif h_level == 2:
                hs.font.size = Pt(14)
            elif h_level == 3:
                hs.font.size = Pt(12)
            else:
                hs.font.size = Pt(11)

        blocks = _parse_markdown(md_content)

        for block in blocks:
            if block.block_type.startswith("h"):
                level = int(block.block_type[1])
                level = min(level, 4)
                heading = doc.add_heading(block.text, level=level)
                for run in heading.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                continue

            if block.block_type in ("ul", "ol"):
                for item_text in block.items:
                    p = doc.add_paragraph(style="List Bullet" if block.block_type == "ul" else "List Number")
                    p.clear()
                    segments = _split_bold_segments(item_text)
                    for seg_text, is_bold in segments:
                        run = p.add_run(seg_text)
                        run.font.name = "微软雅黑"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                        run.font.size = Pt(10.5)
                        run.bold = is_bold
                continue

            if block.block_type == "p":
                p = doc.add_paragraph()
                segments = _split_bold_segments(block.text)
                for seg_text, is_bold in segments:
                    run = p.add_run(seg_text)
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    run.font.size = Pt(10.5)
                    run.bold = is_bold
                continue

        doc.save(output_path)
        print(f"[exporter] DOCX 已导出: {output_path}")

    def to_pdf(self, md_content: str, output_path: str):
        self._init_cn_font()

        font_name = "CNFont" if "CNFont" in pdfmetrics._fonts else "Helvetica"
        font_bold = "CNFontBold" if "CNFontBold" in pdfmetrics._fonts else font_name

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()

        def _make_style(name: str, parent: str, **kwargs) -> ParagraphStyle:
            base = styles[parent]
            return ParagraphStyle(name, parent=base, fontName=font_name, **kwargs)

        style_h1 = _make_style("CNH1", "Heading1", fontSize=16, spaceAfter=8, textColor="black")
        style_h2 = _make_style("CNH2", "Heading2", fontSize=14, spaceAfter=6, textColor="black")
        style_h3 = _make_style("CNH3", "Heading3", fontSize=12, spaceAfter=4, textColor="black")
        style_h4 = _make_style("CNH4", "Heading4", fontSize=11, spaceAfter=4, textColor="black")
        style_body = _make_style("CNBody", "Normal", fontSize=10.5, leading=16, spaceAfter=4)
        style_list = _make_style("CNList", "Normal", fontSize=10.5, leading=16, spaceAfter=2, leftIndent=16)

        heading_styles = {"h1": style_h1, "h2": style_h2, "h3": style_h3, "h4": style_h4}

        story = []
        blocks = _parse_markdown(md_content)

        for block in blocks:
            if block.block_type.startswith("h"):
                level = block.block_type
                hs = heading_styles.get(level, style_h3)
                story.append(Paragraph(_escape_xml(block.text), hs))
                story.append(Spacer(1, 4))
                continue

            if block.block_type in ("ul", "ol"):
                for item_text in block.items:
                    bullet = "•" if block.block_type == "ul" else f"{block.items.index(item_text) + 1}."
                    html = _build_inline_html(item_text, font_name, font_bold)
                    story.append(Paragraph(f"{bullet} {html}", style_list))
                continue

            if block.block_type == "p":
                html = _build_inline_html(block.text, font_name, font_bold)
                story.append(Paragraph(html, style_body))
                continue

        try:
            doc.build(story)
            print(f"[exporter] PDF 已导出: {output_path}")
        except Exception as e:
            print(f"[exporter] PDF 构建失败: {e}")
            raise


def _escape_xml(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_inline_html(text: str, font_name: str, font_bold: str) -> str:
    segments = _split_bold_segments(text)
    parts = []
    for seg_text, is_bold in segments:
        escaped = _escape_xml(seg_text)
        if is_bold:
            parts.append(f'<font face="{font_bold}"><b>{escaped}</b></font>')
        else:
            parts.append(f'<font face="{font_name}">{escaped}</font>')
    return "".join(parts)


def export_resume(md_content: str, output_path: str, fmt: str = "docx"):
    exporter = ResumeExporter()
    fmt = fmt.lower().lstrip(".")

    if fmt == "docx":
        if not output_path.endswith(".docx"):
            output_path += ".docx"
        exporter.to_docx(md_content, output_path)
    elif fmt == "pdf":
        if not output_path.endswith(".pdf"):
            output_path += ".pdf"
        exporter.to_pdf(md_content, output_path)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}，可选: docx, pdf")

    return output_path
